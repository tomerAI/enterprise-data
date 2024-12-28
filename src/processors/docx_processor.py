from typing import List, Dict, Any, Optional
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.section import CT_SectPr
from .base_processor import BaseProcessor
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from datetime import datetime

class DocxProcessor(BaseProcessor):
    PROCESS_TYPES = {
        'default': {'chunk_size': 1000, 'chunk_overlap': 200},
        'fine': {'chunk_size': 500, 'chunk_overlap': 100},
        'coarse': {'chunk_size': 2000, 'chunk_overlap': 400},
        'paragraph': {'chunk_size': None, 'chunk_overlap': 0},
        'page': {'chunk_size': None, 'chunk_overlap': 0}
    }

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.supported_extensions = {'.docx'}
        self.configure_splitter(chunk_size, chunk_overlap)

    def configure_splitter(self, chunk_size: Optional[int], chunk_overlap: int):
        if chunk_size:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        else:
            self.text_splitter = None

    def process(self, file_path: str, process_type: Optional[str] = None) -> List[Dict[str, Any]]:
        """Process a DOCX file and return a list of document chunks with metadata"""
        if process_type and process_type not in self.get_supported_process_types():
            raise ValueError(f"Unsupported process type: {process_type}")
        
        # Configure processor based on process_type
        if process_type:
            config = self.PROCESS_TYPES[process_type]
            self.configure_splitter(config['chunk_size'], config['chunk_overlap'])

        doc = Document(file_path)
        metadata = self._extract_metadata(doc, file_path)
        
        if process_type == 'paragraph':
            return self._process_by_paragraphs(doc, metadata)
        
        return self._process_with_chunking(doc, metadata)

    def _process_by_paragraphs(self, doc: Document, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process document paragraph by paragraph"""
        processed_chunks = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_index": i,
                    "total_chunks": len(doc.paragraphs),
                    "process_type": "paragraph"
                })
                processed_chunks.append({
                    "content": para.text,
                    "metadata": chunk_metadata
                })
        return processed_chunks

    def _process_with_chunking(self, doc: Document, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process document using text chunking"""
        # Combine text from paragraphs and tables
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
            if table_text:
                full_text.append("\n".join(table_text))
        
        # Join all text and create chunks
        combined_text = "\n\n".join(full_text)
        chunks = self.text_splitter.split_text(combined_text)
        
        # Create final chunk objects with metadata
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks)
            })
            
            processed_chunks.append({
                "content": chunk,
                "metadata": chunk_metadata
            })
            
        return processed_chunks
    
    def _extract_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """Extract comprehensive metadata from the document"""
        return {
            "type": "docx",
            "source": file_path,
            "filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "created_at": datetime.now().isoformat(),
            "document_stats": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            },
            "document_properties": {
                "author": doc.core_properties.author or "Unknown",
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                "title": doc.core_properties.title or os.path.splitext(os.path.basename(file_path))[0],
                "subject": doc.core_properties.subject or None,
                "keywords": doc.core_properties.keywords or None,
                "category": doc.core_properties.category or None,
                "comments": doc.core_properties.comments or None
            }
        }

    @classmethod
    def get_supported_extensions(cls) -> set:
        return {'.docx'}

    @classmethod
    def get_supported_process_types(cls) -> set:
        return set(cls.PROCESS_TYPES.keys())

    def _is_page_break(self, paragraph: Paragraph) -> bool:
        """Check if paragraph contains a page break"""
        for run in paragraph._element.findall('.//w:br[@w:type="page"]', 
                                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            return True
        return False

    def _process_by_pages(self, doc: Document, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process document page by page using actual page breaks"""
        processed_chunks = []
        current_page = []
        page_number = 1

        for para in doc.paragraphs:
            # Add paragraph content to current page
            if para.text.strip():
                current_page.append(para.text)

            # Check for page break
            if self._is_page_break(para) or para._element.getparent().tag.endswith('sectPr'):
                if current_page:  # Only process if we have content
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "page_number": page_number,
                        "process_type": "page",
                        "chunk_type": "complete_page"
                    })
                    
                    processed_chunks.append({
                        "content": "\n".join(current_page),
                        "metadata": chunk_metadata
                    })
                    
                    current_page = []
                    page_number += 1

        # Handle any remaining content
        if current_page:
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "page_number": page_number,
                "process_type": "page",
                "chunk_type": "final_page"
            })
            
            processed_chunks.append({
                "content": "\n".join(current_page),
                "metadata": chunk_metadata
            })

        return processed_chunks